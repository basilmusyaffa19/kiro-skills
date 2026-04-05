#!/usr/bin/env python3
# /// script
# requires-python = ">=3.12"
# dependencies = [
#   "pypandoc>=1.17",
#   "python-docx>=1.1.0",
#   "docxcompose>=2.1.0",
# ]
# ///
"""
Markdown to DOCX Converter
Reads Obsidian-flavored Markdown files, preprocesses them, and converts to DOCX via Pandoc.

Bundled with the technical-document-writer skill.
Paths resolve relative to CWD (user's project root), except --template which
defaults to the bundled template inside the skill's assets/ folder.
"""

import argparse
import os
import re
import sys
import tempfile
from pathlib import Path

import pypandoc
from docxcompose.composer import Composer
from docx import Document

from preprocessor import preprocess

# Skill root: two levels up from scripts/md_to_docx.py
SKILL_DIR = Path(__file__).resolve().parent.parent
DEFAULT_TEMPLATE = str(SKILL_DIR / "assets" / "template.docx")


def parse_cover_metadata(filepath: Path) -> dict | None:
    """Parse YAML frontmatter from a cover.md file. Returns dict or None."""
    content = filepath.read_text(encoding="utf-8")
    match = re.match(r"\A---\n(.*?)\n---", content, flags=re.DOTALL)
    if not match:
        return None

    metadata = {}
    current_list_key = None
    current_item = {}

    def _flush_item():
        nonlocal current_item
        if current_list_key and current_item:
            metadata.setdefault(current_list_key, []).append(current_item)
        current_item = {}

    for line in match.group(1).split("\n"):
        if line.strip().startswith("#") or not line.strip():
            _flush_item()
            current_list_key = None
            continue

        if re.match(r"^\s+-\s+\w+:", line):
            _flush_item()
            current_item = {}
            key, _, value = line.strip().lstrip("- ").partition(":")
            current_item[key.strip()] = value.strip().strip('"').strip("'")
            continue

        if re.match(r"^\s{4,}\w+:", line) and current_list_key:
            key, _, value = line.strip().partition(":")
            current_item[key.strip()] = value.strip().strip('"').strip("'")
            continue

        if re.match(r"^\w[\w_]+:\s*$", line):
            _flush_item()
            current_list_key = line.strip().rstrip(":")
            continue

        if ":" in line and not line.startswith(" "):
            _flush_item()
            current_list_key = None
            key, _, value = line.partition(":")
            metadata[key.strip()] = value.strip().strip('"').strip("'")

    _flush_item()
    return metadata


def update_cover_page(doc: Document, metadata: dict) -> None:
    """Replace placeholder text in cover page with metadata values."""
    for paragraph in doc.paragraphs:
        full_text = paragraph.text.strip()
        if full_text == "[…]" and metadata.get("subtitle"):
            _replace_paragraph_text(paragraph, metadata["subtitle"])
        elif full_text == "(…)" and metadata.get("topic"):
            _replace_paragraph_text(paragraph, metadata["topic"])

    for table in doc.tables:
        header_row = table.rows[0]
        header_text = " ".join(c.text.strip() for c in header_row.cells)

        if len(table.rows) == 1 and len(table.columns) == 1:
            cell_text = header_row.cells[0].text.strip()
            if cell_text == "Technical Documentation" and metadata.get("title"):
                _set_cell_text(header_row.cells[0], metadata["title"])
            continue

        if len(table.columns) == 3:
            for row in table.rows:
                label = row.cells[0].text.strip().lower()
                ref_run = None
                if row.cells[2].paragraphs and row.cells[2].paragraphs[0].runs:
                    ref_run = row.cells[2].paragraphs[0].runs[0]

                if label == "author" and metadata.get("author"):
                    _set_cell_text_with_ref(row.cells[1], metadata["author"], ref_run)
                    _set_cell_text(row.cells[2], "")
                elif label == "approved by" and metadata.get("approved_by"):
                    _set_cell_text_with_ref(row.cells[1], metadata["approved_by"], ref_run)
                    _set_cell_text(row.cells[2], "")
                elif label == "version" and metadata.get("version"):
                    _set_cell_text_with_ref(row.cells[1], metadata["version"], ref_run)
                    _set_cell_text(row.cells[2], "")
            continue

        if len(table.columns) == 2 and len(table.rows) >= 5:
            first_label = table.rows[0].cells[0].text.strip().lower()
            if first_label == "document title":
                _update_overview_table(table, metadata)
                continue

        if "VERSION" in header_text and "DESCRIPTION" in header_text:
            _update_revision_table(table, metadata)
            continue

        if "Name" in header_text and "Squad" in header_text:
            _update_authors_table(table, metadata)
            continue


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    """Replace all runs in a paragraph with new text, preserving first run formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = new_text


def _update_overview_table(table, metadata: dict) -> None:
    """Update Document Overview table."""
    field_map = {
        "document title": "document_title",
        "version number": "version_number",
        "reviewer": "reviewer",
        "author": "doc_author",
        "last update": "last_update",
        "status": "status",
    }
    for row in table.rows:
        label = row.cells[0].text.strip().lower()
        meta_key = field_map.get(label)
        if meta_key and metadata.get(meta_key):
            _set_cell_text(row.cells[1], metadata[meta_key])


def _update_revision_table(table, metadata: dict) -> None:
    """Update Revision History table."""
    revisions = metadata.get("revision_history", [])
    for i, rev in enumerate(revisions):
        row_idx = i + 1
        if row_idx < len(table.rows):
            row = table.rows[row_idx]
            if len(row.cells) >= 5:
                _set_cell_text(row.cells[0], rev.get("version", ""))
                _set_cell_text(row.cells[1], rev.get("date", ""))
                _set_cell_text(row.cells[2], rev.get("description", ""))
                _set_cell_text(row.cells[3], rev.get("rev_author", ""))
                _set_cell_text(row.cells[4], rev.get("rev_approved_by", ""))


def _update_authors_table(table, metadata: dict) -> None:
    """Update Authors table."""
    authors = metadata.get("authors", [])
    for i, author in enumerate(authors):
        row_idx = i + 1
        if row_idx < len(table.rows):
            row = table.rows[row_idx]
            if len(row.cells) >= 2:
                _set_cell_text(row.cells[0], author.get("name", ""))
                _set_cell_text(row.cells[1], author.get("squad", ""))


def _set_cell_text(cell, text: str) -> None:
    """Set cell text while preserving formatting of the first run."""
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)

    para = cell.paragraphs[0]

    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        run = para.add_run(text)
        if para.style and para.style.font:
            if para.style.font.name:
                run.font.name = para.style.font.name
            if para.style.font.size:
                run.font.size = para.style.font.size


def _set_cell_text_with_ref(cell, text: str, ref_run=None) -> None:
    """Set cell text, copying font formatting from a reference run."""
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)

    para = cell.paragraphs[0]

    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        run = para.add_run(text)
        if ref_run and ref_run.font:
            if ref_run.font.name:
                run.font.name = ref_run.font.name
            if ref_run.font.size:
                run.font.size = ref_run.font.size
            if ref_run.font.bold is not None:
                run.font.bold = ref_run.font.bold
            if ref_run.font.color and ref_run.font.color.rgb:
                run.font.color.rgb = ref_run.font.color.rgb


def collect_markdown_files(
    source_dir: Path, skip_first: bool = False
) -> tuple[list[Path], Path | None]:
    """Collect and sort markdown files. Returns (content_files, cover_file)."""
    all_files = sorted(
        [source_dir / f for f in os.listdir(source_dir) if f.endswith(".md")]
    )

    if not all_files:
        print(f"No markdown files found in {source_dir}")
        sys.exit(1)

    cover_file = None
    content_files = []
    for f in all_files:
        if "cover" in f.stem.lower():
            cover_file = f
            print(f"Cover metadata: {f.name}")
        elif skip_first and f == all_files[0] and f != cover_file:
            print(f"Skipping: {f.name} (first file)")
        else:
            content_files.append(f)

    if not content_files:
        print("No content markdown files found")
        sys.exit(1)

    print(f"Found {len(content_files)} content file(s) to process")
    return content_files, cover_file


def build_combined_markdown(files: list[Path]) -> str:
    """Read, preprocess, and combine markdown files with page breaks."""
    sections = []

    for filepath in files:
        print(f"Processing: {filepath.name}")
        raw = filepath.read_text(encoding="utf-8")
        processed = preprocess(raw)
        sections.append(processed)

    page_break = '\n\n```{=openxml}\n<w:p><w:r><w:br w:type="page"/></w:r></w:p>\n```\n\n'

    combined = page_break.join(sections)
    return page_break + combined


def convert_to_docx(
    markdown_content: str,
    output_path: Path,
    template_path: Path | None = None,
    toc: bool = True,
    toc_depth: int = 3,
) -> None:
    """Convert markdown string to DOCX using pypandoc."""
    extra_args = []

    if template_path and template_path.exists():
        extra_args.append(f"--reference-doc={template_path}")
        print(f"Using template: {template_path}")

    if toc:
        extra_args.append("--toc")
        extra_args.append(f"--toc-depth={toc_depth}")

    extra_args.append("--wrap=none")

    pypandoc.convert_text(
        markdown_content,
        "docx",
        format="markdown",
        outputfile=str(output_path),
        extra_args=extra_args,
    )

    print(f"\nDocument generated: {output_path}")


def merge_cover_page(
    template_path: Path, content_path: Path, output_path: Path,
    cover_metadata: dict | None = None,
) -> None:
    """Merge cover page from template with Pandoc-generated content."""
    master = Document(str(template_path))

    if cover_metadata:
        update_cover_page(master, cover_metadata)
        print("Cover page updated with metadata")

    composer = Composer(master)
    content_doc = Document(str(content_path))
    composer.append(content_doc)
    composer.save(str(output_path))
    print(f"Cover page merged from: {template_path}")


def apply_justify(docx_path: Path, cover_table_count: int = 0) -> None:
    """Post-process: set all Normal paragraphs to justify alignment and add table borders."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(str(docx_path))

    for paragraph in doc.paragraphs:
        if paragraph.style.name in ("Normal", "Body Text", "First Paragraph"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for i, table in enumerate(doc.tables):
        if i >= cover_table_count:
            _apply_table_borders(table)

    doc.save(str(docx_path))
    print("Post-processing applied (justify + table borders)")


def _apply_table_borders(table) -> None:
    """Add borders to all cells in a table."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "999999")
        borders.append(element)

    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.addprevious(tblPr)


def parse_args() -> argparse.Namespace:
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(
        description="Convert Obsidian Markdown files to DOCX"
    )
    parser.add_argument(
        "--source", "-s",
        default="draft",
        help="Source directory containing .md files (default: draft)",
    )
    parser.add_argument(
        "--template", "-t",
        default=DEFAULT_TEMPLATE,
        help=f"Reference DOCX template (default: bundled assets/template.docx)",
    )
    parser.add_argument(
        "--output", "-o",
        default="output.docx",
        help="Output DOCX file path (default: output.docx)",
    )
    parser.add_argument(
        "--no-toc",
        action="store_true",
        help="Disable Table of Contents generation",
    )
    parser.add_argument(
        "--toc-depth",
        type=int,
        default=3,
        help="TOC depth level (default: 3)",
    )
    parser.add_argument(
        "--skip-first",
        action="store_true",
        help="Skip the first markdown file in sort order",
    )
    return parser.parse_args()


def main() -> None:
    """Main entry point. Paths resolve relative to CWD."""
    args = parse_args()

    source_dir = Path(args.source).resolve()
    template_path = Path(args.template).resolve()
    output_path = Path(args.output).resolve()

    if not source_dir.exists():
        print(f"Source directory not found: {source_dir}")
        sys.exit(1)

    files, cover_file = collect_markdown_files(source_dir, skip_first=args.skip_first)
    combined = build_combined_markdown(files)

    cover_metadata = None
    if cover_file:
        cover_metadata = parse_cover_metadata(cover_file)

    has_template = template_path.exists()

    if has_template:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            temp_path = Path(tmp.name)

        convert_to_docx(
            markdown_content=combined,
            output_path=temp_path,
            template_path=template_path,
            toc=not args.no_toc,
            toc_depth=args.toc_depth,
        )

        merge_cover_page(template_path, temp_path, output_path, cover_metadata)

        template_doc = Document(str(template_path))
        cover_table_count = len(template_doc.tables)

        apply_justify(output_path, cover_table_count)
        temp_path.unlink()
    else:
        convert_to_docx(
            markdown_content=combined,
            output_path=output_path,
            template_path=None,
            toc=not args.no_toc,
            toc_depth=args.toc_depth,
        )
        apply_justify(output_path)


if __name__ == "__main__":
    main()
